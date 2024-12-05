from flask import Flask, render_template, request, send_file
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from docx import Document
from docx.shared import Pt, RGBColor
from docx2pdf import convert
import io
import os
# import pypandoc
# pypandoc.download_pandoc()

app = Flask(__name__)

@app.route('/', methods=['GET', 'POST'])
@app.route('/', methods=['GET', 'POST'])

def report_form():
    if request.method == 'POST':
        # Get form data
        ar_no = request.form['ar_no']
        date = request.form['date']
        product = request.form['product']
        analysed_on = request.form['analysed_on']
        purchaser = request.form['purchaser']
        batch_no = request.form['batch_no']
        po_no = request.form['po_no']
        po_date = request.form['po_date']
        qty_ordered = request.form['qty_ordered']
        qty_dispatched = request.form['qty_dispatched']

         # Load the Word template
        doc = Document('./data/Sample_Analysis Report_KashyapVapi.docx')

        for p in doc.paragraphs:
            print(p.text)
            replace_placeholders(p, ar_no, date, analysed_on, purchaser, batch_no, po_no, po_date, qty_ordered, qty_dispatched)
        
        # Replace placeholders in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        replace_placeholders(paragraph, ar_no, date, analysed_on, purchaser, batch_no, po_no, po_date, qty_ordered, qty_dispatched)

        # Create output directory if it doesn't exist
        output_dir = os.path.join(os.getcwd(), 'output')
        os.makedirs(output_dir, exist_ok=True)

        # # Save the document to the output folder
        # output_path = os.path.join(output_dir, 'analysis_report.docx')
        # doc.save(output_path)

        # Save the document to the output folder as DOCX
        docx_output_path = os.path.join(output_dir, 'analysis_report.docx')
        doc.save(docx_output_path)

        # Convert the DOCX file to PDF and save it in the output folder
        pdf_output_path = os.path.join(output_dir, 'analysis_report.pdf')
        convert(docx_output_path, pdf_output_path)

        # Serve the file for download
        return send_file(pdf_output_path, as_attachment=True, download_name='analysis_report.pdf', mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        
        # # Save the modified document to a BytesIO object
        # doc_io = io.BytesIO()
        # doc.save(doc_io)
        # doc_io.seek(0)
        # return send_file(doc_io, as_attachment=True, download_name='analysis_report.docx', mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

        # # Optionally convert to PDF here if needed
        # # pdf_io = convert_docx_to_pdf(doc_io)

        # # # Save the Word document to a temporary file
        # temp_docx_path = '/Users/Naman/Documents/naman/crm/output/tmp/temp_analysis_report.docx'
        # with open(temp_docx_path, 'wb') as f:
        #     f.write(doc_io.getvalue())

        # # Convert the Word document to PDF using pypandoc
        # output_pdf_path = '/Users/Naman/Documents/naman/crm/output/tmp/analysis_report.pdf'
        # # pypandoc.convert_file(temp_docx_path, 'pdf', outputfile=output_pdf_path)
        # convert(temp_docx_path, output_pdf_path)

        # return send_file(output_pdf_path, as_attachment=True, download_name='analysis_report.pdf', mimetype='application/pdf')
        

        
    
    return render_template('report_form_v2.html')

# def replace_placeholders(p, ar_no, date, analysed_on, purchaser, batch_no, po_no, po_date, qty_ordered, qty_dispatched):
#         if '{{A.R_No}}' in p.text:
#                 p.text = p.text.replace('{{A.R_No}}', ar_no)
#             # if '{{Product}}' in p.text:
#             #     p.text = p.text.replace('{{Product}}', product)
#         if '{{invoice_date}}' in p.text:
#             p.text = p.text.replace('{{invoice_date}}', date)
#         if '{{Analysed_On}}' in p.text:
#             p.text = p.text.replace('{{Analysed_On}}', analysed_on)
#         if '{{Purchaser}}' in p.text:
#             p.text = p.text.replace('{{Purchaser}}', purchaser)
#         if '{{Batch_No}}' in p.text:
#             p.text = p.text.replace('{{Batch_No}}', batch_no)
#         if '{{P.O_No}}' in p.text:
#             p.text = p.text.replace('{{P.O_No}}', po_no)
#         if '{{P.O_Date}}' in p.text:
#             p.text = p.text.replace('{{P.O_Date}}', po_date)
#         if '{{Qty_Ordered}}' in p.text:
#             p.text = p.text.replace('{{Qty_Ordered}}', f"{qty_ordered} KG")
#         if '{{Qty_Dispatched}}' in p.text:
#             p.text = p.text.replace('{{Qty_Dispatched}}', f"{qty_dispatched} KG")

# def replace_placeholders(paragraph, ar_no, date, analysed_on, purchaser, batch_no, po_no, po_date, qty_ordered, qty_dispatched):
#     # Replace placeholders and set font style
#     replacements = {
#         '{{A.R_No}}': ar_no,
#         '{{invoice_date}}': date,
#         '{{Analysed_On}}': analysed_on,
#         '{{Purchaser}}': purchaser,
#         '{{Batch_No}}': batch_no,
#         '{{P.O_No}}': po_no,
#         '{{P.O_Date}}': po_date,
#         '{{Qty_Ordered}}': f"{qty_ordered} KG",
#         '{{Qty_Dispatched}}': f"{qty_dispatched} KG"
#     }

#     for placeholder, value in replacements.items():
#         if placeholder in paragraph.text:
#             # # Clear the paragraph text
#             # paragraph.clear()
#             paragraph.text = paragraph.text.replace(placeholder, value)
#             # Add the new text with the desired style
#             # run = paragraph.add_run(value)
#             # run.font.name = 'Times New Roman'
#             # run.font.size = Pt(14)  # Set font size
#             # run.font.color.rgb = RGBColor(0, 0, 0)  # Set font color (black)
#             # You can set other styles as needed



def replace_placeholders(paragraph, ar_no, date, analysed_on, purchaser, batch_no, po_no, po_date, qty_ordered, qty_dispatched):
    # Replace placeholders and set font style
    replacements = {
        '{{A.R_No}}': ar_no,
        '{{invoice_date}}': date,
        '{{Analysed_On}}': analysed_on,
        '{{Purchaser}}': purchaser,
        '{{Batch_No}}': batch_no,
        '{{P.O_No}}': po_no,
        '{{P.O_Date}}': po_date,
        '{{Qty_Ordered}}': f"{qty_ordered} KG",
        '{{Qty_Dispatched}}': f"{qty_dispatched} KG"
    }

    for placeholder, value in replacements.items():
        if placeholder in paragraph.text:
            # Replace placeholder in paragraph text
            paragraph.text = paragraph.text.replace(placeholder, value)

            # Rebuild runs to style the replaced value
            new_runs = paragraph.text.split(value)
            paragraph.clear()  # Clear all current runs

            # Add runs back with formatting for replaced text
            for idx, run_text in enumerate(new_runs):
                if run_text:  # Add unstyled part
                    paragraph.add_run(run_text)

                if idx < len(new_runs) - 1:  # Add styled replacement
                    styled_run = paragraph.add_run(value)
                    styled_run.font.name = 'Times New Roman'
                    styled_run.font.size = Pt(14)
                    styled_run.font.color.rgb = RGBColor(0, 0, 0)



if __name__ == '__main__':
    app.run(debug=True)
